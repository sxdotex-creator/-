#!/usr/bin/env python3
"""Excel -> Word 模板填充工具

根据 Excel 表格中的列值，自动替换 Word 模板中的占位符，
并生成多份 Word 文档。

支持跨平台（Windows / macOS / Linux），仅依赖 openpyxl 和 python-docx。
"""

import argparse
import json
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import openpyxl
from docx import Document


PLACEHOLDER_PREFIX = "{{"
PLACEHOLDER_SUFFIX = "}}"


def normalize_header(value: Optional[str]) -> str:
    if value is None:
        return ""
    return re.sub(r"[\s_\-]+", "", str(value)).strip().lower()


def parse_map_args(map_args: List[str]) -> Dict[str, str]:
    mappings: Dict[str, str] = {}
    for item in map_args:
        if "=" not in item:
            raise ValueError(f"映射参数格式错误，应为 header=PLACEHOLDER: {item}")
        source, target = item.split("=", 1)
        source = source.strip()
        target = target.strip()
        if not source or not target:
            raise ValueError(f"映射参数不能为空: {item}")
        mappings[source] = target
    return mappings


def load_json_config(config_path: str) -> Dict:
    with open(config_path, "r", encoding="utf-8") as fp:
        return json.load(fp)


def find_header_index(headers: List[str], name: str) -> int:
    norm_name = normalize_header(name)
    for index, header in enumerate(headers):
        if normalize_header(header) == norm_name:
            return index
    raise KeyError(f"Excel 表头未找到: {name}")


def load_excel_rows(
    excel_path: str,
    sheet_name: Optional[str] = None,
    start_row: int = 2,
) -> Tuple[List[str], List[List[str]]]:
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    worksheet = workbook[sheet_name] if sheet_name else workbook.active

    all_rows = list(worksheet.iter_rows(values_only=True))
    if not all_rows:
        raise ValueError("Excel 文件中没有数据")

    headers = [str(cell).strip() if cell is not None else "" for cell in all_rows[0]]
    body = [list(row) for row in all_rows[start_row - 1 :]]
    workbook.close()
    return headers, body


def format_docx_placeholder(key: str, prefix: str = PLACEHOLDER_PREFIX, suffix: str = PLACEHOLDER_SUFFIX) -> str:
    return f"{prefix}{key}{suffix}"


def normalize_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def replace_placeholders_in_text(text: str, values: Dict[str, str], prefix: str, suffix: str) -> str:
    for key, value in values.items():
        placeholder = format_docx_placeholder(key, prefix, suffix)
        text = text.replace(placeholder, value)
    return text


def replace_placeholders_in_paragraph(paragraph, values: Dict[str, str], prefix: str, suffix: str) -> None:
    for run in paragraph.runs:
        new_text = replace_placeholders_in_text(run.text, values, prefix, suffix)
        if new_text != run.text:
            run.text = new_text


def replace_placeholders_in_table(table, values: Dict[str, str], prefix: str, suffix: str) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values, prefix, suffix)


def replace_placeholders_in_document(doc: Document, values: Dict[str, str], prefix: str, suffix: str) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, values, prefix, suffix)
    for table in doc.tables:
        replace_placeholders_in_table(table, values, prefix, suffix)
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values, prefix, suffix)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values, prefix, suffix)


def build_values_from_row(
    headers: List[str],
    row: List,
    mapping: Dict[str, str],
    prefix: str,
    suffix: str,
    row_index: int,
) -> Dict[str, str]:
    values: Dict[str, str] = {
        "ROW_INDEX": str(row_index),
        "ROW_NUMBER": str(row_index),
    }
    for excel_header, placeholder_name in mapping.items():
        column_index = find_header_index(headers, excel_header)
        values[placeholder_name] = normalize_value(row[column_index] if column_index < len(row) else "")
    return values


def build_output_file_name(pattern: str, values: Dict[str, str]) -> str:
    try:
        return pattern.format(**values)
    except KeyError as exc:
        raise ValueError(f"输出文件名模板缺少占位符: {exc}") from exc


def validate_template_file(path: str) -> None:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Word 模板文件不存在: {path}")
    if not path.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 格式的 Word 模板")


def validate_excel_file(path: str) -> None:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Excel 文件不存在: {path}")


def create_output_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def build_default_output_pattern(headers: List[str], mapping: Dict[str, str]) -> str:
    if "BOX_NO" in mapping.values():
        return "{BOX_NO}.docx"
    if "ID" in mapping.values() or "SEQ" in mapping.values():
        return "{ROW_INDEX}.docx"
    return "record_{ROW_INDEX}.docx"


def run(args: argparse.Namespace) -> None:
    mapping = {}
    output_pattern = args.output_name or "record_{ROW_INDEX}.docx"
    prefix = PLACEHOLDER_PREFIX
    suffix = PLACEHOLDER_SUFFIX

    if args.config:
        config = load_json_config(args.config)
        mapping = config.get("mappings", {})
        output_pattern = config.get("output_name", output_pattern)
        placeholder_style = config.get("placeholder_style", {})
        prefix = placeholder_style.get("prefix", prefix)
        suffix = placeholder_style.get("suffix", suffix)
    else:
        mapping = parse_map_args(args.map)

    if not mapping:
        raise ValueError("必须通过 --map 或 --config 提供列映射关系")

    validate_excel_file(args.excel)
    validate_template_file(args.template)
    create_output_dir(args.output)

    headers, rows = load_excel_rows(args.excel, sheet_name=args.sheet, start_row=args.start_row)
    if not rows:
        raise ValueError("Excel 文件中没有可用数据行")

    if args.output_name is None and not args.config:
        output_pattern = build_default_output_pattern(headers, mapping)

    print(f"读取到 {len(rows)} 行数据，正在生成文档...")

    seen_names = set()
    for index, row in enumerate(rows, start=1):
        if all(cell is None or str(cell).strip() == "" for cell in row):
            continue

        values = build_values_from_row(headers, row, mapping, prefix, suffix, row_index=index)
        doc = Document(args.template)
        replace_placeholders_in_document(doc, values, prefix, suffix)

        raw_name = build_output_file_name(output_pattern, values)
        safe_name = Path(raw_name).stem
        extension = Path(raw_name).suffix or ".docx"
        output_name = f"{safe_name}{extension}"
        if output_name in seen_names:
            output_name = f"{safe_name}_{index}{extension}"
        seen_names.add(output_name)

        output_path = Path(args.output) / output_name
        doc.save(output_path)
        print(f"  • 已生成：{output_path}")

    print("\n全部文档生成完成！")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="根据 Excel 数据填充 Word 模板，并输出多份文档。"
    )
    parser.add_argument("--excel", required=True, help="Excel 文件路径 (.xlsx)")
    parser.add_argument("--template", required=True, help="Word 模板路径 (.docx)")
    parser.add_argument("--output", required=True, help="输出目录")
    parser.add_argument("--sheet", default=None, help="要读取的工作表名称（默认第一个）")
    parser.add_argument("--start-row", type=int, default=2, help="数据起始行，默认 2（跳过表头）")
    parser.add_argument(
        "--map",
        action="append",
        default=[],
        help="列映射: Excel表头=模板占位符，例如 --map 盒号=BOX_NO --map 文号=DOC_NO",
    )
    parser.add_argument(
        "--config",
        default=None,
        help="JSON 配置文件路径，优先于 --map。配置示例见 README。",
    )
    parser.add_argument(
        "--output-name",
        default=None,
        help="输出文件名模板，例如 {BOX_NO}_{DOC_NO}.docx。默认使用 ROW_INDEX。",
    )
    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    try:
        run(args)
    except Exception as exc:
        print(f"错误: {exc}")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
