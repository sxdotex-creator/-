"""核心文档生成逻辑。"""

import copy
import json
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import openpyxl
from docx import Document

DEFAULT_PLACEHOLDER_STYLES = [
    ("{{", "}}"),
    ("<<", ">>"),
    ("[[", "]]")
]
REPEAT_MARKER_KEY = "#repeat"


def normalize_header(value: Optional[str]) -> str:
    if value is None:
        return ""
    return re.sub(r"[\s_\-]+", "", str(value)).strip().lower()


def load_json_config(config_path: str) -> Dict:
    with open(config_path, "r", encoding="utf-8") as fp:
        return json.load(fp)


def parse_map_args(map_args: List[str]) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for item in map_args:
        if "=" not in item:
            raise ValueError(f"映射参数格式错误，应为 header=PLACEHOLDER: {item}")
        source, target = item.split("=", 1)
        source = source.strip()
        target = target.strip()
        if not source or not target:
            raise ValueError(f"映射参数不能为空: {item}")
        mapping[source] = target
    return mapping


def find_header_index(headers: List[str], name: str) -> int:
    norm_name = normalize_header(name)
    for index, header in enumerate(headers):
        if normalize_header(header) == norm_name:
            return index
    raise KeyError(f"未找到 Excel 表头: {name}")


def load_excel_rows(
    excel_path: str,
    sheet_name: Optional[str] = None,
    start_row: int = 2,
) -> Tuple[List[str], List[List[Optional[str]]]]:
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    worksheet = workbook[sheet_name] if sheet_name else workbook.active

    rows = list(worksheet.iter_rows(values_only=True))
    if not rows:
        raise ValueError("Excel 文件中没有数据")

    headers = [str(cell).strip() if cell is not None else "" for cell in rows[0]]
    body = [list(row) for row in rows[start_row - 1 :]]
    workbook.close()
    return headers, body


def format_placeholders(key: str, styles: List[Tuple[str, str]]) -> List[str]:
    return [f"{prefix}{key}{suffix}" for prefix, suffix in styles]


def normalize_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def build_values_from_row(
    headers: List[str],
    row: List[Optional[str]],
    mapping: Dict[str, str],
    row_index: int,
) -> Dict[str, str]:
    values: Dict[str, str] = {
        "ROW_INDEX": str(row_index),
        "ROW_NUMBER": str(row_index),
    }
    for excel_header, placeholder_name in mapping.items():
        column_index = find_header_index(headers, excel_header)
        values[placeholder_name] = normalize_value(row[column_index] if column_index < len(row) else "")
    return values


def replace_placeholders_in_text(
    text: str,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> str:
    if not text:
        return text
    for key, value in values.items():
        for prefix, suffix in placeholder_styles:
            placeholder = f"{prefix}{key}{suffix}"
            if placeholder in text:
                text = text.replace(placeholder, value)
    return text


def replace_placeholders_in_paragraph(
    paragraph,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> None:
    for run in paragraph.runs:
        new_text = replace_placeholders_in_text(run.text, values, placeholder_styles)
        if new_text != run.text:
            run.text = new_text


def replace_placeholders_in_cell(
    cell,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> None:
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, values, placeholder_styles)


def replace_placeholders_in_table(
    table,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            replace_placeholders_in_cell(cell, values, placeholder_styles)


def replace_placeholders_in_row(
    row,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> None:
    for cell in row.cells:
        replace_placeholders_in_cell(cell, values, placeholder_styles)


def replace_placeholders_in_document(
    document: Document,
    values: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
) -> None:
    for paragraph in document.paragraphs:
        replace_placeholders_in_paragraph(paragraph, values, placeholder_styles)
    for table in document.tables:
        replace_placeholders_in_table(table, values, placeholder_styles)
    for section in document.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values, placeholder_styles)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, values, placeholder_styles)


def get_cell_text(cell) -> str:
    return cell.text or ""


def find_repeat_row_index(table, marker_key: str, placeholder_styles: List[Tuple[str, str]]) -> Optional[int]:
    marker_candidates = format_placeholders(marker_key, placeholder_styles)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            text = get_cell_text(cell)
            if any(marker in text for marker in marker_candidates):
                return row_index
    return None


def clone_table_row(table, row_index: int):
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    new_tr = copy.deepcopy(tr)
    tbl.insert(row_index + 1, new_tr)
    return table.rows[row_index + 1]


def remove_table_row(table, row_index: int) -> None:
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    tbl.remove(tr)


def replicate_table_rows(
    document: Document,
    headers: List[str],
    rows: List[List[Optional[str]]],
    mapping: Dict[str, str],
    placeholder_styles: List[Tuple[str, str]],
    marker_key: str,
) -> None:
    for table in document.tables:
        template_index = find_repeat_row_index(table, marker_key, placeholder_styles)
        if template_index is None:
            continue

        start_index = template_index
        for row_index, row in enumerate(rows, start=1):
            values = build_values_from_row(headers, row, mapping, row_index)
            values[marker_key] = ""
            values["GROUP_ROW_COUNT"] = str(len(rows))
            clone = clone_table_row(table, start_index + row_index - 1)
            replace_placeholders_in_row(clone, values, placeholder_styles)

        remove_table_row(table, template_index)


def group_rows_by_header(
    headers: List[str],
    rows: List[List[Optional[str]]],
    group_by: str,
) -> Dict[str, List[List[Optional[str]]]]:
    index = find_header_index(headers, group_by)
    grouped: Dict[str, List[List[Optional[str]]]] = {}
    for row in rows:
        if index < len(row) and row[index] is not None:
            group_key = normalize_value(row[index])
        else:
            group_key = "UNKNOWN"
        grouped.setdefault(group_key, []).append(row)
    return grouped


def build_group_values(
    headers: List[str],
    row: List[Optional[str]],
    mapping: Dict[str, str],
    group_key: str,
    group_size: int,
) -> Dict[str, str]:
    values = build_values_from_row(headers, row, mapping, row_index=1)
    values["GROUP_KEY"] = group_key
    values["GROUP_ROW_COUNT"] = str(group_size)
    return values


def build_output_file_name(pattern: str, values: Dict[str, str]) -> str:
    try:
        return pattern.format(**values)
    except KeyError as exc:
        raise ValueError(f"输出文件名模板缺少占位符: {exc}") from exc


def generate_documents(
    excel_path: str,
    template_path: str,
    output_dir: str,
    mapping: Dict[str, str],
    output_name_pattern: Optional[str] = None,
    sheet_name: Optional[str] = None,
    start_row: int = 2,
    group_by: Optional[str] = None,
    repeat_marker: str = REPEAT_MARKER_KEY,
    placeholder_styles: Optional[List[Tuple[str, str]]] = None,
) -> List[Path]:
    placeholder_styles = placeholder_styles or DEFAULT_PLACEHOLDER_STYLES
    output_name_pattern = output_name_pattern or "record_{ROW_INDEX}.docx"

    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel 文件不存在: {excel_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Word 模板文件不存在: {template_path}")
    if not template_path.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 格式的 Word 模板")

    os.makedirs(output_dir, exist_ok=True)

    headers, rows = load_excel_rows(excel_path, sheet_name=sheet_name, start_row=start_row)
    if not rows:
        raise ValueError("Excel 文件中没有可用数据行")

    result_files: List[Path] = []
    existing_names = set()

    if group_by:
        grouped = group_rows_by_header(headers, rows, group_by)
        for group_key, group_rows in grouped.items():
            group_values = build_group_values(headers, group_rows[0], mapping, group_key, len(group_rows))
            document = Document(template_path)
            replace_placeholders_in_document(document, group_values, placeholder_styles)
            replicate_table_rows(
                document,
                headers,
                group_rows,
                mapping,
                placeholder_styles,
                marker_key=repeat_marker,
            )

            raw_name = build_output_file_name(output_name_pattern, group_values)
            safe_name = Path(raw_name).stem
            ext = Path(raw_name).suffix or ".docx"
            output_name = f"{safe_name}{ext}"
            if output_name in existing_names:
                output_name = f"{safe_name}_{group_key}{ext}"
            existing_names.add(output_name)
            output_path = Path(output_dir) / output_name
            document.save(output_path)
            result_files.append(output_path)
    else:
        for index, row in enumerate(rows, start=1):
            if all(cell is None or str(cell).strip() == "" for cell in row):
                continue
            values = build_values_from_row(headers, row, mapping, index)
            values[REPEAT_MARKER_KEY] = ""
            document = Document(template_path)
            replace_placeholders_in_document(document, values, placeholder_styles)

            raw_name = build_output_file_name(output_name_pattern, values)
            safe_name = Path(raw_name).stem
            ext = Path(raw_name).suffix or ".docx"
            output_name = f"{safe_name}{ext}"
            if output_name in existing_names:
                output_name = f"{safe_name}_{index}{ext}"
            existing_names.add(output_name)
            output_path = Path(output_dir) / output_name
            document.save(output_path)
            result_files.append(output_path)

    return result_files
